"""多格式导出器——支持 MD, TXT, EPUB, PDF, DOCX。

EPUB: 完整 CSS 排版，卷章导航，元数据，封面。
PDF: 专业中文排版，页眉页脚，页码，目录。
"""

from __future__ import annotations
import re
from pathlib import Path
from project import Project

# ──────────────────────────────────────────────
# CSS for EPUB
# ──────────────────────────────────────────────
EPUB_CSS = """
@namespace epub "http://www.idpf.org/2007/ops";

body {
    font-family: "PingFang SC", "Noto Serif CJK SC", "STSong", serif;
    line-height: 1.85;
    margin: 5%;
    text-align: justify;
}

h1 {
    font-size: 1.8em;
    text-align: center;
    margin: 2em 0 1em;
    border-bottom: 2px solid #333;
    padding-bottom: 0.5em;
    page-break-before: always;
}

h1:first-of-type {
    page-break-before: avoid;
}

h2 {
    font-size: 1.4em;
    text-align: center;
    margin: 1.5em 0 0.8em;
}

h3 {
    font-size: 1.2em;
    text-align: center;
    margin: 1em 0 0.5em;
    font-weight: normal;
    color: #555;
}

p {
    text-indent: 2em;
    margin: 0.5em 0;
}

p.no-indent {
    text-indent: 0;
}

.scene-break {
    text-align: center;
    margin: 1.5em 0;
    letter-spacing: 0.5em;
    color: #999;
}

.volume-title {
    font-size: 1.8em;
    text-align: center;
    font-weight: bold;
    margin: 2em 0 0.5em;
    page-break-before: always;
}

.volume-synopsis {
    font-style: italic;
    text-align: center;
    color: #666;
    margin-bottom: 1.5em;
}

.title-page {
    text-align: center;
    margin-top: 15%;
}

.title-page h1 {
    font-size: 2.2em;
    border: none;
    margin-bottom: 0.5em;
}

.title-page .author {
    font-size: 1.2em;
    color: #666;
    margin: 1em 0;
}

.title-page .genre {
    font-size: 1em;
    color: #999;
}
"""


class Exporter:
    """将项目导出为多种格式。"""

    def __init__(self, project: Project, output_dir: str = ""):
        self.project = project
        self.output_dir = Path(output_dir) if output_dir else (
            Path.home() / "Desktop" / "DeepSeekWriter" / project.title
        )

    def export_all(self) -> dict[str, str]:
        results = {}
        results["md"] = self.export_markdown()
        results["txt"] = self.export_txt()
        results["epub"] = self.export_epub()
        results["pdf"] = self.export_pdf()
        results["docx"] = self.export_docx()
        return results

    # ═══════════════════════════════════════════
    # Audio (m4a / mp3 via edge-tts)
    # ═══════════════════════════════════════════

    def export_audio(self, fmt: str = "m4a", per_chapter: bool = True,
                     voice: str = "") -> str:
        """导出有声书。

        Args:
            fmt: "m4a" 或 "mp3"
            per_chapter: True=每章单独文件, False=整本合并
            voice: TTS 声优（空=默认）

        Returns: 输出目录路径
        """
        from utils.audio import AudioExporter
        exporter = AudioExporter(self.project, str(self.output_dir), voice=voice or "")
        if fmt == "mp3":
            return exporter.export_mp3(per_chapter)
        return exporter.export_m4a(per_chapter)

    # ═══════════════════════════════════════════
    # Markdown
    # ═══════════════════════════════════════════

    def export_markdown(self) -> str:
        self.output_dir.mkdir(parents=True, exist_ok=True)
        path = self.output_dir / f"{self.project.title}.md"
        lines = [
            f"# 《{self.project.title}》",
            "",
            f"> **类型**: {self.project.genre}　**主题**: {self.project.theme}　**基调**: {self.project.tone}",
            f"> {self.project.premise}",
            "",
        ]
        # Table of contents
        lines.append("## 目录")
        lines.append("")
        for vol in self.project.volumes:
            lines.append(f"- **第{vol.volume_number}卷「{vol.volume_title}」**")
            for ch in vol.chapters:
                status = "✅" if ch.content else "⏳"
                lines.append(f"  - {status} 第{ch.chapter_number}章「{ch.chapter_title}」")
        lines.append("")

        for vol in self.project.volumes:
            lines.append(f"# 第{vol.volume_number}卷「{vol.volume_title}」")
            lines.append(f"> {vol.synopsis}")
            lines.append("")
            for ch in vol.chapters:
                if ch.content:
                    lines.append(f"## 第{ch.chapter_number}章「{ch.chapter_title}」")
                    lines.append("")
                    lines.append(ch.content)
                    lines.append("")
                    lines.append("---")
                    lines.append("")
        path.write_text("\n".join(lines), encoding="utf-8")
        return str(path)

    # ═══════════════════════════════════════════
    # TXT
    # ═══════════════════════════════════════════

    def export_txt(self) -> str:
        self.output_dir.mkdir(parents=True, exist_ok=True)
        path = self.output_dir / f"{self.project.title}.txt"
        lines = [
            f"《{self.project.title}》",
            f"类型: {self.project.genre}  |  主题: {self.project.theme}  |  基调: {self.project.tone}",
            f"{self.project.premise}",
            "",
            "═" * 50,
            "",
        ]
        for vol in self.project.volumes:
            lines.append(f"第{vol.volume_number}卷「{vol.volume_title}」")
            lines.append(f"  {vol.synopsis}")
            lines.append("─" * 40)
            for ch in vol.chapters:
                if ch.content:
                    lines.append(f"第{ch.chapter_number}章「{ch.chapter_title}」")
                    lines.append("")
                    lines.append(ch.content)
                    lines.append("")
                    lines.append("─" * 40)
                    lines.append("")
        path.write_text("\n".join(lines), encoding="utf-8")
        return str(path)

    # ═══════════════════════════════════════════
    # EPUB — fully formatted
    # ═══════════════════════════════════════════

    def export_epub(self) -> str:
        self.output_dir.mkdir(parents=True, exist_ok=True)
        path = self.output_dir / f"{self.project.title}.epub"
        try:
            self._write_epub(str(path))
            return str(path)
        except ImportError:
            return "EPUB 导出需要: pip install ebooklib"

    def _write_epub(self, path: str):
        from ebooklib import epub

        book = epub.EpubBook()
        book.set_title(self.project.title)
        book.set_language("zh")
        book.add_author("DeepSeek Writer")
        book.add_metadata("DC", "description", self.project.premise[:500])
        book.add_metadata("DC", "subject", self.project.genre)
        book.add_metadata("DC", "publisher", "DeepSeek Writer (AI-Assisted)")
        book.add_metadata("DC", "date", self.project.created_at[:10] if self.project.created_at else "")

        # CSS
        style = epub.EpubItem(
            uid="style", file_name="style/default.css",
            media_type="text/css", content=EPUB_CSS,
        )
        book.add_item(style)

        # ── Title page ──
        title_page = epub.EpubHtml(
            title="扉页", file_name="title_page.xhtml", lang="zh",
        )
        title_page.content = f"""<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml" xml:lang="zh">
<head><meta charset="utf-8"/><link rel="stylesheet" type="text/css" href="style/default.css"/></head>
<body>
<div class="title-page">
<h1>《{self.project.title}》</h1>
<div class="author">DeepSeek Writer</div>
<div class="genre">{self.project.genre}　·　{self.project.theme}</div>
<p class="no-indent" style="margin-top:2em;">{self.project.premise}</p>
</div>
</body></html>"""
        title_page.add_item(style)
        book.add_item(title_page)

        spine = ["nav", title_page]
        toc = []

        # ── Character page ──
        chars = self.project.characters.get("characters", [])
        if chars:
            char_page = epub.EpubHtml(
                title="人物表", file_name="characters.xhtml", lang="zh",
            )
            char_html = """<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml" xml:lang="zh">
<head><meta charset="utf-8"/><link rel="stylesheet" type="text/css" href="style/default.css"/></head>
<body>
<h1>人物表</h1>"""
            for c in chars:
                char_html += f"""
<h2>{c.get('name','')}（{c.get('role','')}）</h2>
<p class="no-indent"><strong>年龄:</strong> {c.get('age','')}　<strong>外貌:</strong> {c.get('appearance','')}</p>
<p class="no-indent"><strong>性格:</strong> {c.get('personality','')}</p>
<p class="no-indent"><strong>动机:</strong> {c.get('motivation','')}</p>
<p class="no-indent"><strong>弧线:</strong> {c.get('arc','')}</p>"""
            char_html += "\n</body></html>"
            char_page.content = char_html
            char_page.add_item(style)
            book.add_item(char_page)
            spine.append(char_page)

        # ── Volumes & Chapters ──
        for vol in self.project.volumes:
            vn = vol.volume_number
            vt = vol.volume_title

            # Volume intro page
            has_content = any(ch.content for ch in vol.chapters)
            if has_content:
                vol_page = epub.EpubHtml(
                    title=f"第{vn}卷「{vt}」",
                    file_name=f"vol_{vn:02d}.xhtml",
                    lang="zh",
                )
                vol_page.content = f"""<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml" xml:lang="zh">
<head><meta charset="utf-8"/><link rel="stylesheet" type="text/css" href="style/default.css"/></head>
<body>
<div class="volume-title">第{vn}卷「{vt}」</div>
<div class="volume-synopsis">{vol.synopsis}</div>
</body></html>"""
                vol_page.add_item(style)
                book.add_item(vol_page)
                spine.append(vol_page)

            for ch in vol.chapters:
                if not ch.content:
                    continue
                ch_html = epub.EpubHtml(
                    title=ch.chapter_title,
                    file_name=f"v{vn:02d}_c{ch.chapter_number:03d}.xhtml",
                    lang="zh",
                )
                # Format content: handle scene breaks and paragraphs
                formatted = self._format_html_content(ch.content)
                ch_html.content = f"""<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml" xml:lang="zh">
<head><meta charset="utf-8"/><link rel="stylesheet" type="text/css" href="style/default.css"/></head>
<body>
<h2>第{ch.chapter_number}章「{ch.chapter_title}」</h2>
{formatted}
</body></html>"""
                ch_html.add_item(style)
                book.add_item(ch_html)
                spine.append(ch_html)
                toc.append(ch_html)

        # ── TOC ──
        book.toc = [epub.Link("title_page.xhtml", "扉页", "title")]
        if chars:
            book.toc.append(epub.Link("characters.xhtml", "人物表", "characters"))

        for vol in self.project.volumes:
            vol_items = []
            for ch in vol.chapters:
                if ch.content:
                    vol_items.append(epub.Link(
                        f"v{vol.volume_number:02d}_c{ch.chapter_number:03d}.xhtml",
                        f"第{ch.chapter_number}章「{ch.chapter_title}」",
                        f"v{vol.volume_number}c{ch.chapter_number}",
                    ))
            if vol_items:
                book.toc.append((
                    epub.Section(f"第{vol.volume_number}卷「{vol.volume_title}」"),
                    tuple(vol_items),
                ))

        book.spine = spine
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        epub.write_epub(path, book)

    def _format_html_content(self, content: str) -> str:
        """将小说正文转为 HTML 段落，保留场景分隔符。"""
        paragraphs = content.split("\n")
        out = []
        for p in paragraphs:
            p = p.strip()
            if not p:
                out.append("<p>&#160;</p>")
                continue
            # Scene breaks
            if re.match(r'^[\*\-]{3,}$', p) or p in ("***", "---", "* * *"):
                out.append('<p class="scene-break">* * *</p>')
                continue
            # Escape HTML
            p = p.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            out.append(f"<p>{p}</p>")
        return "\n".join(out)

    # ═══════════════════════════════════════════
    # PDF — professional Chinese typesetting
    # ═══════════════════════════════════════════

    def export_pdf(self) -> str:
        self.output_dir.mkdir(parents=True, exist_ok=True)
        path = self.output_dir / f"{self.project.title}.pdf"
        try:
            self._write_pdf(str(path))
            return str(path)
        except ImportError:
            return "PDF 导出需要: pip install fpdf2"

    def _write_pdf(self, path: str):
        from fpdf import FPDF

        pdf = FPDF(orientation="P", unit="mm", format="A5")
        # Add Chinese font
        font_paths = [
            "/System/Library/Fonts/PingFang.ttc",
            "/System/Library/Fonts/STHeiti Light.ttc",
            "/System/Library/Fonts/Hiragino Sans GB.ttc",
            "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        ]
        font_loaded = False
        for fp in font_paths:
            if Path(fp).exists():
                pdf.add_font("CJK", "", fp, uni=True)
                font_loaded = True
                break
        if not font_loaded:
            # Fallback: try to find any CJK font
            for fp in Path("/System/Library/Fonts").glob("*.ttc"):
                try:
                    pdf.add_font("CJK", "", str(fp), uni=True)
                    font_loaded = True
                    break
                except Exception:
                    continue
        if not font_loaded:
            raise RuntimeError("未找到中文字体，PDF 导出需要中文字体文件")

        pdf.set_auto_page_break(auto=True, margin=18)

        # ── Title page ──
        pdf.add_page()
        pdf.ln(40)
        pdf.set_font("CJK", "", 24)
        pdf.multi_cell(0, 14, f"《{self.project.title}》", align="C")
        pdf.ln(8)
        pdf.set_font("CJK", "", 12)
        pdf.cell(0, 8, f"{self.project.genre}　·　{self.project.theme}", align="C", new_x="LMARGIN", new_y="NEXT")
        pdf.cell(0, 8, f"基调: {self.project.tone}", align="C", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(10)
        pdf.set_font("CJK", "", 10)
        pdf.multi_cell(0, 7, self.project.premise, align="C")
        pdf.ln(10)
        pdf.set_font("CJK", "", 9)
        pdf.cell(0, 6, f"创作于 {self.project.created_at[:10] if self.project.created_at else ''}", align="C", new_x="LMARGIN", new_y="NEXT")
        pdf.cell(0, 6, "DeepSeek Writer (AI-Assisted)", align="C", new_x="LMARGIN", new_y="NEXT")

        # ── Character page ──
        chars = self.project.characters.get("characters", [])
        if chars:
            pdf.add_page()
            pdf.set_font("CJK", "", 18)
            pdf.cell(0, 12, "人物表", align="C", new_x="LMARGIN", new_y="NEXT")
            pdf.ln(6)
            for c in chars:
                # Check if we need a new page
                if pdf.y > 160:
                    pdf.add_page()
                pdf.set_font("CJK", "", 13)
                pdf.cell(0, 8, f"{c.get('name','')}（{c.get('role','')}）", new_x="LMARGIN", new_y="NEXT")
                pdf.set_font("CJK", "", 9)
                pdf.multi_cell(0, 6,
                    f"年龄: {c.get('age','')}　外貌: {c.get('appearance','')}\n"
                    f"性格: {c.get('personality','')}\n"
                    f"动机: {c.get('motivation','')}",
                )
                pdf.ln(2)

        # ── Chapters ──
        for vol in self.project.volumes:
            vn = vol.volume_number
            vt = vol.volume_title

            # Volume title page
            pdf.add_page()
            pdf.ln(30)
            pdf.set_font("CJK", "", 22)
            pdf.multi_cell(0, 12, f"第{vn}卷「{vt}」", align="C")
            pdf.ln(6)
            pdf.set_font("CJK", "", 10)
            pdf.multi_cell(0, 7, vol.synopsis, align="C")

            for ch in vol.chapters:
                if not ch.content:
                    continue
                pdf.add_page()
                # Chapter header
                pdf.set_font("CJK", "", 16)
                pdf.cell(0, 10, f"第{ch.chapter_number}章「{ch.chapter_title}」", align="C", new_x="LMARGIN", new_y="NEXT")
                pdf.ln(4)

                # Chapter body
                pdf.set_font("CJK", "", 10)
                for line in ch.content.split("\n"):
                    line = line.strip()
                    if not line:
                        pdf.ln(4)
                        continue
                    # Scene break
                    if re.match(r'^[\*\-]{3,}$', line) or line in ("***", "---", "* * *"):
                        pdf.set_font("CJK", "", 9)
                        pdf.cell(0, 6, "＊　＊　＊", align="C", new_x="LMARGIN", new_y="NEXT")
                        pdf.set_font("CJK", "", 10)
                        pdf.ln(2)
                        continue
                    # Normal paragraph with 2-em indent
                    pdf.set_x(pdf.l_margin + 20)
                    pdf.multi_cell(
                        pdf.w - pdf.l_margin - pdf.r_margin - 20,
                        6.5, line, align="J",
                    )
                    pdf.ln(0.5)

                # Page number at bottom
                pdf.set_font("CJK", "", 8)
                pdf.set_y(-15)
                pdf.cell(0, 8, f"— {pdf.page_no()} —", align="C")

        pdf.output(path)

    # ═══════════════════════════════════════════
    # DOCX
    # ═══════════════════════════════════════════

    def export_docx(self) -> str:
        self.output_dir.mkdir(parents=True, exist_ok=True)
        path = self.output_dir / f"{self.project.title}.docx"
        try:
            self._write_docx(str(path))
            return str(path)
        except ImportError:
            return "DOCX 导出需要: pip install python-docx"

    def _write_docx(self, path: str):
        from docx import Document
        from docx.shared import Pt, Cm, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn

        doc = Document()

        # Default style
        style = doc.styles["Normal"]
        font = style.font
        font.name = "PingFang SC"
        font.size = Pt(11)
        style.paragraph_format.line_spacing = 1.6
        style.paragraph_format.space_after = Pt(6)
        # Set East Asian font
        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rFonts.set(qn("w:eastAsia"), "PingFang SC")
        rPr.insert(0, rFonts)

        # ── Title ──
        title = doc.add_heading(f"《{self.project.title}》", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ── Info ──
        info = doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info.add_run(f"类型: {self.project.genre}　主题: {self.project.theme}　基调: {self.project.tone}").italic = True
        doc.add_paragraph(self.project.premise).alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

        # ── Character table ──
        chars = self.project.characters.get("characters", [])
        if chars:
            doc.add_heading("人物表", 1)
            for c in chars:
                doc.add_heading(f"{c.get('name','')}（{c.get('role','')}）", 2)
                p = doc.add_paragraph()
                p.add_run(f"年龄: {c.get('age','')}").bold = True
                doc.add_paragraph(f"外貌: {c.get('appearance','')}")
                doc.add_paragraph(f"性格: {c.get('personality','')}")
                doc.add_paragraph(f"动机: {c.get('motivation','')}")
                doc.add_paragraph(f"弧线: {c.get('arc','')}")
            doc.add_page_break()

        # ── Volumes & Chapters ──
        for vol in self.project.volumes:
            doc.add_heading(f"第{vol.volume_number}卷「{vol.volume_title}」", 1)
            syn = doc.add_paragraph(vol.synopsis)
            syn.italic = True

            for ch in vol.chapters:
                if not ch.content:
                    continue
                doc.add_heading(f"第{ch.chapter_number}章「{ch.chapter_title}」", 2)
                for line in ch.content.split("\n"):
                    line = line.strip()
                    if not line:
                        doc.add_paragraph("")
                    elif re.match(r'^[\*\-]{3,}$', line) or line in ("***", "---", "* * *"):
                        sep = doc.add_paragraph("＊　＊　＊")
                        sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        p = doc.add_paragraph(line)
                        p.paragraph_format.first_line_indent = Cm(0.74)  # 2 Chinese chars

        doc.save(path)
